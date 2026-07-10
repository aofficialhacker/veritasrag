"""
build_blackbook_v2.py
=====================
Builds the VeritasRAG blackbook by REUSING the exact front-matter pages from the
previous semester's journal (title page, proforma, certificate, acknowledgement,
declaration) and splicing in the new VeritasRAG content (abstract, table of
contents, chapters, references, appendices).

Only the project title and the author name are changed on the reused pages.

Usage:
    python build_blackbook_v2.py "Vishal Rajeshkumar Gor" VeritasRAG_Blackbook_Vishal.docx
    python build_blackbook_v2.py "Nayanta Shinde"          VeritasRAG_Blackbook_Nayanta.docx
"""
import os
import sys

from docx import Document
from docx.oxml.ns import qn

import build_blackbook as bb  # reuse content builders + helpers

HERE = os.path.dirname(__file__)
PREV_DOCX = r"D:\FREELANCING PROJECT 2\RAG PROJECT\Real_Time_Stock_Monitor_Documentation vishal.docx"
RESEARCH_PAPER_PDF = r"D:\FREELANCING PROJECT 2\RAG PROJECT\RAG_Research_Paper.pdf"

OLD_TITLE = "Real-Time Stock Monitoring with Kafka"
OLD_AUTHOR = "Vishal Rajeshkumar Gor"


# ---------------------------------------------------------------------------
def find_heading(doc, text):
    """Return the Paragraph whose text matches `text` (case-insensitive)."""
    target = text.strip().upper()
    for p in doc.paragraphs:
        if p.text.strip().upper() == target:
            return p
    return None


def delete_range(start_elem, end_elem):
    """Delete sibling elements from start_elem up to (not including) end_elem."""
    cur = start_elem
    while cur is not None and cur is not end_elem:
        nxt = cur.getnext()
        cur.getparent().remove(cur)
        cur = nxt


def replace_in_paragraph(p, mapping):
    full = p.text
    new = full
    for k, v in mapping.items():
        if k in new:
            new = new.replace(k, v)
    if new != full:
        if p.runs:
            for r in p.runs[1:]:
                r.text = ""
            p.runs[0].text = new
        else:
            p.add_run(new)


def replace_everywhere(doc, mapping):
    for p in doc.paragraphs:
        replace_in_paragraph(p, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, mapping)


# ---------------------------------------------------------------------------
def build(author, out_path):
    doc = Document(PREV_DOCX)
    body = doc.element.body

    abstract_h = find_heading(doc, "ABSTRACT")
    ack_h = find_heading(doc, "ACKNOWLEDGEMENT")
    toc_h = find_heading(doc, "TABLE OF CONTENTS")
    sectPr = body.find(qn("w:sectPr"))

    if not (abstract_h and ack_h and toc_h):
        raise RuntimeError("Could not locate expected front-matter headings in previous doc.")

    # 1. Remove the OLD (Kafka) abstract: from ABSTRACT heading up to ACKNOWLEDGEMENT.
    delete_range(abstract_h._p, ack_h._p)

    # 2. Remove everything from TABLE OF CONTENTS to the end (old TOC, lists, chapters).
    delete_range(toc_h._p, sectPr)

    # 3. Build the NEW VeritasRAG abstract, then move it to sit before ACKNOWLEDGEMENT.
    before = list(body)
    bb.abstract_page(doc)
    new_elems = [e for e in list(body) if e not in before]
    for e in new_elems:
        body.remove(e)
        ack_h._p.addprevious(e)

    # 4. Append the rest of the VeritasRAG content after the (kept) declaration.
    bb.toc_pages(doc)
    bb.chapter1(doc)
    bb.chapter2(doc)
    bb.chapter3(doc)
    bb.chapter4(doc)
    bb.chapter5(doc)
    bb.chapter6(doc)
    bb.chapter7(doc)
    bb.chapter8(doc)
    bb.chapter9(doc)
    bb.references(doc)
    bb.appendices(doc)

    # 5. Attach the research paper (rendered pages) after the appendices.
    if os.path.exists(RESEARCH_PAPER_PDF):
        bb.attach_pdf_as_images(doc, RESEARCH_PAPER_PDF, "ANNEXURE: RESEARCH PAPER")

    # 6. Swap the project title and author on the reused front-matter pages.
    replace_everywhere(doc, {OLD_TITLE: bb.TITLE, OLD_AUTHOR: author})

    # 7. Put a black box border on every page.
    bb.add_page_borders(doc)

    bb.enable_update_fields(doc)
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    author = sys.argv[1] if len(sys.argv) > 1 else OLD_AUTHOR
    out = sys.argv[2] if len(sys.argv) > 2 else "VeritasRAG_Blackbook.docx"
    build(author, os.path.join(HERE, out))
